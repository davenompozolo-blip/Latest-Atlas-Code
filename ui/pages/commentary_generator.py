"""
ATLAS Terminal - AI-Powered Investment Commentary Generator
============================================================
Generates structured investment commentary from live ATLAS data using Claude.
Supports three commentary types:
  1. Quarterly Positioning & Outlook
  2. Manager Research Summary (5Ps framework)
  3. Portfolio Attribution Commentary

Exports to .docx with RisCura-appropriate formatting.

Created: February 2026
Author: Hlobo & Claude
"""
from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Optional

import streamlit as st

from ui.theme import ATLAS_COLORS as THEME


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_COMMENTARY_TYPES = [
    "Quarterly Positioning & Outlook",
    "Manager Research Summary",
    "Portfolio Attribution Commentary",
]

_TONE_OPTIONS = ["Institutional", "Concise", "Detailed"]

_WORD_COUNTS = {
    "Institutional": 800,
    "Concise": 400,
    "Detailed": 1200,
}

# Header colour — reads from branding config
try:
    from config.branding import get_docx_config as _get_docx
    _HEADER_RGB = tuple(_get_docx()["header_rgb"])
except Exception:
    _HEADER_RGB = (31, 41, 63)

# ---------------------------------------------------------------------------
# System prompts
# ---------------------------------------------------------------------------

_PROMPT_QUARTERLY = """\
You are a senior investment strategist writing for a South African institutional \
investment consultancy. Your commentary is read by pension fund trustees, CIOs, \
and investment committees. The tone is authoritative, precise, and grounded in \
evidence. You use complete sentences. You do not use bullet points unless \
specifically structured as a Q&A. You cite specific data points where provided. \
You acknowledge uncertainty without hedging excessively.

Write a Quarterly Positioning & Outlook commentary based on the following data \
inputs. Structure it with these sections:

1. Market Backdrop — What happened in the period. Reference specific return \
numbers and regime context where provided.
2. Portfolio Positioning — How the portfolio was positioned and why. Link \
positioning to the macro view.
3. Key Calls — The 2-3 highest-conviction views and the rationale behind each.
4. Outlook — Forward-looking view for the next quarter. Be specific about risks \
and opportunities.

Do not pad. Target {word_count} words.\
"""

_PROMPT_MANAGER = """\
You are a fund manager research analyst at a South African institutional \
investment consultancy. You write structured manager assessments following the \
5Ps framework: Philosophy, Process, People, Performance, and Portfolio. Your \
audience is investment committees who use these assessments to make hire/fire/\
watch decisions.

Write a Manager Research Summary using the 5Ps framework based on the data \
inputs below. Each section should be 2-3 sentences. Be specific. Flag any \
concerns or areas requiring further due diligence. Conclude with an overall \
assessment (Recommended / Watch / Not Recommended) with a one-sentence \
justification.

Do not pad. Target {word_count} words.\
"""

_PROMPT_ATTRIBUTION = """\
You are a performance analyst at a South African institutional investment \
consultancy. You write attribution commentaries that explain what drove \
portfolio returns in plain institutional language. Your audience is CIOs and \
portfolio managers who need a narrative interpretation of the numbers.

Write a Portfolio Attribution Commentary for {period} based on the data inputs \
below. Structure it as:

1. Headline — One sentence summarising the period's outcome (absolute and \
relative return).
2. What Drove Returns — Decompose into allocation effect and selection effect. \
Name the top contributors and detractors specifically.
3. What Changed — Any notable position changes or rebalancing during the period.
4. Looking Forward — One paragraph on implications for next period.

Do not pad. Target {word_count} words.\
"""


# ---------------------------------------------------------------------------
# Claude API call
# ---------------------------------------------------------------------------

def _generate_commentary(system_prompt: str, user_message: str) -> str:
    """Call Anthropic API and return the generated commentary text."""
    import os
    import anthropic

    api_key = st.secrets.get("anthropic", {}).get("api_key", "")
    if not api_key:
        api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise ValueError(
            "Anthropic API key not configured. Set ANTHROPIC_API_KEY environment "
            "variable or add [anthropic] api_key to .streamlit/secrets.toml"
        )

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )
    return response.content[0].text.strip()


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def _build_docx(commentary: str, commentary_type: str, date_str: str) -> bytes:
    """Build a .docx file with RisCura-appropriate formatting.

    Returns raw bytes suitable for st.download_button.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # -- Modify built-in styles for Calibri --
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.15

    style_h2 = doc.styles["Heading 2"]
    style_h2.font.name = "Calibri"
    style_h2.font.size = Pt(13)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(*_HEADER_RGB)

    # -- Header bar --
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        from config.branding import get_branding as _gb
        _firm = _gb()["firm_name"]
    except Exception:
        _firm = "ATLAS Terminal"
    run = header_para.add_run(f"{_firm}  —  {commentary_type}")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*_HEADER_RGB)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(date_str)
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()  # spacer

    # -- Body: parse sections by numbered headings or newlines --
    lines = commentary.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Detect section headings (e.g., "1. Market Backdrop", "## Heading")
        is_heading = False
        heading_text = stripped
        if stripped and stripped[0].isdigit() and ". " in stripped[:5]:
            heading_text = stripped.split(". ", 1)[1] if ". " in stripped else stripped
            is_heading = True
        elif stripped.startswith("## "):
            heading_text = stripped[3:]
            is_heading = True
        elif stripped.startswith("**") and stripped.endswith("**"):
            heading_text = stripped.strip("*").strip()
            is_heading = True

        if is_heading:
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(stripped)

    # -- Footer --
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"{date_str}  |  {commentary_type}  |  ATLAS Terminal")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(160, 160, 160)

    # -- Serialise --
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Data assembly helpers
# ---------------------------------------------------------------------------

def _build_quarterly_message(
    regime_ctx: Optional[dict],
    market_returns: str,
    portfolio_perf: str,
    key_calls: str,
) -> str:
    """Assemble the user message for Quarterly Positioning."""
    lines = ["Market Backdrop:"]
    if regime_ctx:
        quant = regime_ctx.get("quant_regime", "unknown")
        pm = regime_ctx.get("pm_regime", "unknown")
        gv = regime_ctx.get("growth_value", "unknown")
        consensus = regime_ctx.get("consensus", "unknown")
        lines.append(f"- Current regime: Quantitative={quant}, Returns-based={pm} (Consensus: {consensus})")
        lines.append(f"- Growth/Value tilt: {gv}")
    else:
        lines.append("- Regime context not available")

    if market_returns.strip():
        lines.append("")
        lines.append("Market Returns:")
        lines.append(market_returns)

    if portfolio_perf.strip():
        lines.append("")
        lines.append("Portfolio Performance:")
        lines.append(portfolio_perf)

    if key_calls.strip():
        lines.append("")
        lines.append("Key Calls (user-specified):")
        for i, call in enumerate(key_calls.strip().split("\n")[:3], 1):
            lines.append(f"{i}. {call.strip()}")

    return "\n".join(lines)


def _build_manager_message(
    fund_data: Optional[dict],
    tenure: str,
    aum: str,
    mandate_type: str,
    fee_structure: str,
) -> str:
    """Assemble the user message for Manager Research Summary."""
    lines = ["Manager Research Inputs:"]

    if fund_data:
        lines.append(f"- Fund: {fund_data.get('fund_name', 'Unknown')}")
        lines.append(f"- Ticker: {fund_data.get('manager_name', 'Unknown')}")
        lines.append(f"- Asset class: {fund_data.get('asset_class', 'Unknown')}")
        lines.append(f"- Expense ratio: {fund_data.get('expense_ratio', 'N/A')}")
        perf = fund_data.get('performance_summary', {})
        if perf:
            lines.append(f"- Annualized return: {perf.get('ann_return', 'N/A')}")
            lines.append(f"- Annualized volatility: {perf.get('ann_vol', 'N/A')}")
            lines.append(f"- Sharpe ratio: {perf.get('sharpe', 'N/A')}")
            lines.append(f"- Max drawdown: {perf.get('max_drawdown', 'N/A')}")
    else:
        lines.append("- Fund data not available from Fund Research module")

    lines.append("")
    lines.append("Additional Context:")
    lines.append(f"- Manager tenure: {tenure or 'Not specified'}")
    lines.append(f"- AUM: {aum or 'Not specified'}")
    lines.append(f"- Mandate type: {mandate_type or 'Not specified'}")
    lines.append(f"- Fee structure: {fee_structure or 'Not specified'}")

    return "\n".join(lines)


def _build_attribution_message(
    attr_data: Optional[dict],
    period_label: str,
    benchmark_name: str,
) -> str:
    """Assemble the user message for Portfolio Attribution Commentary."""
    lines = ["Attribution Data:"]

    if attr_data:
        lines.append(f"- Period: {attr_data.get('period', period_label)}")
        lines.append(f"- Benchmark: {attr_data.get('benchmark_name', benchmark_name)}")
        lines.append(f"- Portfolio return: {attr_data.get('total_return', 0):.2f}%")
        lines.append(f"- Benchmark return: {attr_data.get('benchmark_return', 0):.2f}%")
        excess = attr_data.get('total_return', 0) - attr_data.get('benchmark_return', 0)
        lines.append(f"- Excess return: {excess:+.2f}%")
        lines.append(f"- Allocation effect: {attr_data.get('allocation_effect', 0):.2f}%")
        lines.append(f"- Selection effect: {attr_data.get('selection_effect', 0):.2f}%")
        te = attr_data.get('tracking_error', 0)
        ir = attr_data.get('information_ratio', 0)
        if te:
            lines.append(f"- Tracking error: {te:.2f}%")
        if ir:
            lines.append(f"- Information ratio: {ir:.2f}")
        contribs = attr_data.get('top_contributors', [])
        if contribs:
            lines.append(f"- Top contributors: {', '.join(str(c) for c in contribs)}")
        detractors = attr_data.get('top_detractors', [])
        if detractors:
            lines.append(f"- Top detractors: {', '.join(str(d) for d in detractors)}")
    else:
        lines.append("- Attribution data not available from Performance Suite")
        lines.append(f"- Period: {period_label}")
        lines.append(f"- Benchmark: {benchmark_name}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main render function
# ---------------------------------------------------------------------------

def render_commentary_generator():
    """AI-Powered Investment Commentary Generator."""

    st.markdown(
        '<h1 style="font-size: 2.2rem; font-weight: 800;'
        ' color: var(--text-primary, rgba(255,255,255,0.92)); margin-bottom: 0;">'
        '<span style="background: linear-gradient(135deg,'
        f' {THEME["primary"]}, {THEME["secondary"]});'
        ' -webkit-background-clip: text; -webkit-text-fill-color: transparent;">'
        'INVESTMENT COMMENTARY</span></h1>',
        unsafe_allow_html=True,
    )
    st.caption("Draft structured investment commentary from live ATLAS data")

    # ── Commentary type selector ──
    st.markdown("---")
    commentary_type = st.radio(
        "Commentary Type",
        _COMMENTARY_TYPES,
        horizontal=True,
        key="cg_type",
    )

    # ── Tone and word count ──
    col_tone, col_wc = st.columns(2)
    with col_tone:
        tone = st.selectbox("Tone", _TONE_OPTIONS, key="cg_tone")
    with col_wc:
        word_count = st.number_input(
            "Target Word Count",
            min_value=200,
            max_value=2000,
            value=_WORD_COUNTS.get(tone, 800),
            step=100,
            key="cg_word_count",
        )

    st.markdown("---")

    # ── Type-specific inputs ──
    if commentary_type == _COMMENTARY_TYPES[0]:
        system_prompt, user_message = _render_quarterly_inputs(word_count)
    elif commentary_type == _COMMENTARY_TYPES[1]:
        system_prompt, user_message = _render_manager_inputs(word_count)
    else:
        system_prompt, user_message = _render_attribution_inputs(word_count)

    # ── Generate / Regenerate ──
    st.markdown("---")
    col_gen, col_regen = st.columns([3, 1])
    with col_gen:
        generate = st.button(
            "Generate Commentary",
            type="primary",
            use_container_width=True,
            key="cg_generate",
        )
    with col_regen:
        regenerate = st.button(
            "Regenerate",
            use_container_width=True,
            key="cg_regenerate",
        )

    if generate or regenerate:
        with st.spinner("Drafting commentary..."):
            try:
                text = _generate_commentary(system_prompt, user_message)
                st.session_state["commentary_output"] = text
                st.session_state["commentary_type"] = commentary_type
            except Exception as exc:
                st.error(f"Commentary generation error: {exc}")
                return

    # ── Output ──
    output = st.session_state.get("commentary_output")
    if output:
        st.markdown("##### Draft Commentary")
        edited = st.text_area(
            "Edit below",
            value=output,
            height=400,
            key="cg_edited_output",
            label_visibility="collapsed",
        )

        # Word count indicator
        wc = len(edited.split())
        st.caption(f"Word count: {wc}")

        # Download button
        date_str = datetime.now().strftime("%B %Y")
        ct = st.session_state.get("commentary_type", commentary_type)
        try:
            docx_bytes = _build_docx(edited, ct, date_str)
            filename = f"ATLAS_{ct.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
            st.download_button(
                "Download as .docx",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.warning(
                "python-docx is not installed. Run `pip install python-docx` to enable DOCX export."
            )


# ---------------------------------------------------------------------------
# Type-specific input renderers
# ---------------------------------------------------------------------------

def _render_quarterly_inputs(word_count: int) -> tuple[str, str]:
    """Render inputs and return (system_prompt, user_message) for Quarterly Positioning."""

    regime_ctx = st.session_state.get("macro_regime")
    if regime_ctx:
        st.success("Macro regime context detected from Market Regime page.")
    else:
        st.info("No macro regime context in session. Enter manually or run Market Regime first.")

    market_returns = st.text_area(
        "Market Returns (key index returns for the period)",
        placeholder="e.g.\nJSE ALSI: +4.2%\nMSCI World: +6.1%\nALBI: +1.8%",
        height=100,
        key="cg_market_returns",
    )

    portfolio_perf = st.text_area(
        "Portfolio Performance",
        placeholder="e.g.\nAbsolute return: +3.8%\nRelative to benchmark: -0.4%\nPrimary contributors: SA banks, gold miners",
        height=100,
        key="cg_portfolio_perf",
    )

    # Pre-populate from Performance Suite if available
    attr_data = st.session_state.get("attribution_output")
    if attr_data and not portfolio_perf.strip():
        st.caption(
            f"Performance Suite data available: "
            f"portfolio {attr_data.get('total_return', 0):.1f}%, "
            f"benchmark {attr_data.get('benchmark_return', 0):.1f}%"
        )

    key_calls = st.text_area(
        "Key Calls / Themes (max 3 bullet points)",
        placeholder="e.g.\nOverweight SA banks on improving credit cycle\nReduce duration ahead of SARB hiking cycle",
        height=80,
        key="cg_key_calls",
    )

    system_prompt = _PROMPT_QUARTERLY.format(word_count=word_count)
    user_msg = _build_quarterly_message(regime_ctx, market_returns, portfolio_perf, key_calls)
    return system_prompt, user_msg


def _render_manager_inputs(word_count: int) -> tuple[str, str]:
    """Render inputs and return (system_prompt, user_message) for Manager Research."""

    fund_data = st.session_state.get("fund_research_output")
    if fund_data:
        st.success(
            f"Fund Research data detected: {fund_data.get('fund_name', 'Unknown')} "
            f"({fund_data.get('manager_name', '')})"
        )
    else:
        st.info("No fund data in session. Enter manually or run Fund Research first.")

    col1, col2 = st.columns(2)
    with col1:
        tenure = st.text_input("Manager Tenure", placeholder="e.g. 12 years", key="cg_tenure")
        aum = st.text_input("AUM", placeholder="e.g. R 4.2 billion", key="cg_aum")
    with col2:
        mandate_type = st.text_input("Mandate Type", placeholder="e.g. SA Multi-Asset High Equity", key="cg_mandate")
        fee_structure = st.text_input("Fee Structure", placeholder="e.g. 0.85% TER, no performance fee", key="cg_fees")

    system_prompt = _PROMPT_MANAGER.format(word_count=word_count)
    user_msg = _build_manager_message(fund_data, tenure, aum, mandate_type, fee_structure)
    return system_prompt, user_msg


def _render_attribution_inputs(word_count: int) -> tuple[str, str]:
    """Render inputs and return (system_prompt, user_message) for Attribution Commentary."""

    attr_data = st.session_state.get("attribution_output")
    if attr_data:
        st.success(
            f"Attribution data detected from Performance Suite: "
            f"portfolio {attr_data.get('total_return', 0):.1f}%, "
            f"benchmark {attr_data.get('benchmark_return', 0):.1f}%"
        )
    else:
        st.info("No attribution data in session. Enter manually or run Performance Suite first.")

    col1, col2 = st.columns(2)
    with col1:
        period_label = st.text_input(
            "Period Label",
            value="Q4 2025",
            key="cg_period",
        )
    with col2:
        benchmark_name = st.text_input(
            "Benchmark Name",
            value=attr_data.get("benchmark_name", "ALSI") if attr_data else "ALSI",
            key="cg_benchmark",
        )

    system_prompt = _PROMPT_ATTRIBUTION.format(word_count=word_count, period=period_label)
    user_msg = _build_attribution_message(attr_data, period_label, benchmark_name)
    return system_prompt, user_msg
